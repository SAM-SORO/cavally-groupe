"""Extraction multimodale des listes de fournitures via Gemini Flash.

Un seul chemin d'extraction : Word, PDF et images partent tous vers Gemini.
Le .docx est simplement normalise en texte (ou en images s'il n'en contient
que), le PDF et les images sont envoyes tels quels en `inline_data`.
"""

from __future__ import annotations

import io
import json
import logging
import re
import time
from dataclasses import dataclass
from pathlib import Path

from google import genai
from google.genai import types

from .config import get_settings
from .schemas import Extraction

logger = logging.getLogger("cavally.extraction")

# --------------------------------------------------------------------------- #
# Formats acceptes
# --------------------------------------------------------------------------- #

IMAGE_MIME = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
    ".heic": "image/heic",
    ".heif": "image/heif",
}
PDF_MIME = {".pdf": "application/pdf"}
WORD_EXT = {".docx"}
TEXT_EXT = {".txt", ".md"}

SUPPORTED_EXTENSIONS = sorted(set(IMAGE_MIME) | set(PDF_MIME) | WORD_EXT | TEXT_EXT)

DOCX_MAGIC = b"PK\x03\x04"


class ExtractionError(Exception):
    """Erreur fonctionnelle destinee a etre affichee a l'utilisateur."""

    def __init__(self, message: str, status_code: int = 422) -> None:
        super().__init__(message)
        self.message = message
        self.status_code = status_code


# --------------------------------------------------------------------------- #
# Prompt (base : CLAUDE.md, enrichi des metadonnees du modele Excel de reference)
# --------------------------------------------------------------------------- #

PROMPT = """Tu analyses une liste de fournitures scolaires. Le document peut etre
un texte, un PDF ou une image (capture ou photo).

Chaque etablissement redige sa liste a sa maniere : la mise en page, le
vocabulaire, le decoupage en rubriques et le niveau de detail changent d'un
document a l'autre. Tu dois t'adapter au document que tu as sous les yeux, pas
a un modele attendu. Ne transpose jamais dans ta reponse une information qui
n'est pas dans ce document.

Renvoie UNIQUEMENT du JSON valide, sans aucun texte avant ou apres, sans
balises Markdown.

DECOUPAGE PAR CLASSE
- "listes" contient une entree par classe traitee dans le document.
- Un document ne concernant qu'une seule classe donne une seule entree ; un
  fascicule couvrant plusieurs niveaux (CP1, CP2, CE1...) donne une entree par
  niveau, chacune avec ses propres articles.
- Si aucune classe n'est identifiable, mets une seule entree avec "classe": "".

ARTICLES
- "designation" : reprends le libelle tel qu'il apparait, debarrasse des puces
  et numeros de liste. Conserve les precisions utiles qui figurent dans le
  document (format, nombre de pages, marque, couleur, edition, matiere).
- "qte" : le nombre indique. Si aucune quantite n'est precisee, mets 1.
- "categorie" : l'intitule de la rubrique SOUS LAQUELLE l'article est range
  dans le document, repris tel quel (ex. "MANUELS", "TENUE DE SPORT",
  "HYGIENE", "ARTS PLASTIQUES", "TROUSSE"). N'utilise pas une nomenclature
  imposee : c'est le document qui decide. Si le document ne comporte aucune
  rubrique, deduis toi-meme un intitule court (un ou deux mots) et applique le
  meme intitule a tous les articles de meme nature.
- Ignore ce qui n'est pas un article : titres de rubrique, en-tetes, nom de
  l'ecole ou de la classe, mentions generales (ex. "Rentree 2025").
- N'invente rien et ne regroupe pas les articles. Un article = une ligne.
- Conserve l'ordre d'apparition dans le document.
- Si le document est illisible ou ne contient aucune fourniture, renvoie
  "listes": [].

EN-TETE — chaine vide si l'information est absente, n'invente jamais
- "etablissement" : nom de l'ecole / du groupe scolaire.
- "coordonnees" : adresse et/ou telephone de l'etablissement, sur une ligne.
- "annee_scolaire" : annee scolaire, telle qu'ecrite (ex. "2026-2027").
- "classe" (dans chaque entree de "listes") : telle qu'ecrite ("8eme (CM1)").
- "consignes" (dans chaque entree) : uniquement les recommandations explicites
  adressees aux parents (ce qu'il faut faire, marquer, couvrir, apporter),
  une phrase par entree, sans puce, recopiees du document. N'y mets NI les
  titres, NI les sous-titres, NI les mentions de mise en page. Liste VIDE s'il
  n'y en a pas — n'en fabrique aucune."""

_ARTICLE_SCHEMA = types.Schema(
    type=types.Type.OBJECT,
    property_ordering=["categorie", "designation", "qte"],
    required=["categorie", "designation", "qte"],
    properties={
        # Volontairement libre : la rubrique vient du document, pas d'un enum.
        "categorie": types.Schema(type=types.Type.STRING),
        "designation": types.Schema(type=types.Type.STRING),
        "qte": types.Schema(type=types.Type.INTEGER),
    },
)

RESPONSE_SCHEMA = types.Schema(
    type=types.Type.OBJECT,
    property_ordering=["etablissement", "coordonnees", "annee_scolaire", "listes"],
    required=["etablissement", "coordonnees", "annee_scolaire", "listes"],
    properties={
        "etablissement": types.Schema(type=types.Type.STRING),
        "coordonnees": types.Schema(type=types.Type.STRING),
        "annee_scolaire": types.Schema(type=types.Type.STRING),
        "listes": types.Schema(
            type=types.Type.ARRAY,
            items=types.Schema(
                type=types.Type.OBJECT,
                property_ordering=["classe", "consignes", "articles"],
                required=["classe", "consignes", "articles"],
                properties={
                    "classe": types.Schema(type=types.Type.STRING),
                    "consignes": types.Schema(
                        type=types.Type.ARRAY, items=types.Schema(type=types.Type.STRING)
                    ),
                    "articles": types.Schema(type=types.Type.ARRAY, items=_ARTICLE_SCHEMA),
                },
            ),
        ),
    },
)


# --------------------------------------------------------------------------- #
# Normalisation des entrees
# --------------------------------------------------------------------------- #


@dataclass
class DocumentParts:
    """Ce qui part reellement au modele, plus une trace du chemin emprunte."""

    parts: list[types.Part]
    source: str  # "texte" | "pdf" | "image"


def _docx_to_text(payload: bytes) -> str:
    from docx import Document  # import tardif : dependance lourde

    document = Document(io.BytesIO(payload))
    lines: list[str] = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            # Les cellules fusionnees sont repetees par python-docx : on deduplique.
            deduped: list[str] = []
            for cell in cells:
                if cell and (not deduped or deduped[-1] != cell):
                    deduped.append(cell)
            if deduped:
                lines.append(" | ".join(deduped))
    return "\n".join(lines).strip()


def _docx_images(payload: bytes) -> list[tuple[bytes, str]]:
    """Images embarquees dans un .docx (cas d'une capture collee dans Word)."""
    from docx import Document

    document = Document(io.BytesIO(payload))
    images: list[tuple[bytes, str]] = []
    for part in document.part.package.iter_parts():
        content_type = getattr(part, "content_type", "") or ""
        if content_type in IMAGE_MIME.values():
            blob = getattr(part, "blob", b"")
            if blob:
                images.append((blob, content_type))
    return images[:8]


def build_document_parts(filename: str, payload: bytes) -> DocumentParts:
    """Transforme le fichier uploade en `Part` Gemini, sans parseur par format."""
    extension = Path(filename or "").suffix.lower()

    if extension in TEXT_EXT:
        text = payload.decode("utf-8", errors="replace").strip()
        if not text:
            raise ExtractionError("Le fichier texte est vide.")
        return DocumentParts([types.Part.from_text(text=text)], "texte")

    if extension in WORD_EXT:
        if not payload.startswith(DOCX_MAGIC):
            raise ExtractionError(
                "Ce fichier .docx semble corrompu ou il s'agit d'un ancien .doc. "
                "Enregistre-le au format .docx ou exporte-le en PDF."
            )
        try:
            text = _docx_to_text(payload)
        except Exception as exc:  # pragma: no cover - dependant du fichier
            logger.warning("Lecture .docx impossible : %s", exc)
            raise ExtractionError("Impossible de lire ce document Word. Essaie un export PDF.") from exc

        if text:
            return DocumentParts([types.Part.from_text(text=text)], "texte")

        images = _docx_images(payload)
        if images:
            logger.info("docx sans texte : %d image(s) embarquee(s) envoyee(s) au modele", len(images))
            return DocumentParts(
                [types.Part.from_bytes(data=blob, mime_type=mime) for blob, mime in images],
                "image",
            )
        raise ExtractionError("Ce document Word ne contient ni texte ni image exploitable.")

    if extension in PDF_MIME:
        return DocumentParts([types.Part.from_bytes(data=payload, mime_type="application/pdf")], "pdf")

    if extension in IMAGE_MIME:
        return DocumentParts(
            [types.Part.from_bytes(data=payload, mime_type=IMAGE_MIME[extension])],
            "image",
        )

    raise ExtractionError(
        "Format non supporte. Formats acceptes : " + ", ".join(SUPPORTED_EXTENSIONS) + "."
    )


# --------------------------------------------------------------------------- #
# Appel Gemini
# --------------------------------------------------------------------------- #


_client: genai.Client | None = None


def get_client() -> genai.Client:
    global _client
    if _client is None:
        settings = get_settings()
        if not settings.gemini_api_key:
            raise ExtractionError(
                "Cle API Gemini absente. Renseigne GEMINI_API_KEY dans backend/.env.",
                status_code=500,
            )
        _client = genai.Client(api_key=settings.gemini_api_key)
    return _client


def _parse_payload(raw: str) -> dict:
    """Parsing defensif : le modele peut encadrer sa reponse malgre les consignes."""
    text = (raw or "").strip()
    if not text:
        raise ExtractionError("Le modele n'a rien renvoye. Reessaie avec un document plus lisible.")

    fence = re.match(r"^```(?:json)?\s*(.*?)\s*```$", text, flags=re.S)
    if fence:
        text = fence.group(1).strip()

    try:
        data = json.loads(text)
    except json.JSONDecodeError:
        start, end = text.find("{"), text.rfind("}")
        if start == -1 or end <= start:
            start, end = text.find("["), text.rfind("]")
        if start == -1 or end <= start:
            logger.error("Reponse Gemini non parsable : %s", text[:400])
            raise ExtractionError("Reponse du modele illisible. Reessaie dans un instant.")
        try:
            data = json.loads(text[start : end + 1])
        except json.JSONDecodeError as exc:
            logger.error("Reponse Gemini non parsable : %s", text[:400])
            raise ExtractionError("Reponse du modele illisible. Reessaie dans un instant.") from exc

    # Tolerance : le modele peut renvoyer directement le tableau d'articles.
    if isinstance(data, list):
        data = {"listes": [{"classe": "", "consignes": [], "articles": data}]}
    if not isinstance(data, dict):
        raise ExtractionError("Reponse du modele inattendue. Reessaie dans un instant.")
    # Tolerance : forme plate {..., "articles": [...]} sans decoupage par classe.
    if "listes" not in data and isinstance(data.get("articles"), list):
        data = {
            **{c: data.get(c, "") for c in ("etablissement", "coordonnees", "annee_scolaire")},
            "listes": [
                {
                    "classe": data.get("classe", ""),
                    "consignes": data.get("consignes", []),
                    "articles": data["articles"],
                }
            ],
        }
    return data


def _thinking_variants() -> list[types.ThinkingConfig | None]:
    """Facons de brider le raisonnement, de la plus economique a la plus permissive.

    Les Gemini 3.x attendent `thinking_level`, les 2.5 `thinking_budget`. On
    essaie dans l'ordre et on memorise la variante acceptee par le modele
    configure, pour ne payer la detection qu'une seule fois.
    """
    return [
        types.ThinkingConfig(thinking_level="low"),
        types.ThinkingConfig(thinking_budget=0),
        None,
    ]


# Variante retenue, par modele : `gemini-flash-latest` peut basculer sur une
# nouvelle generation qui n'accepte plus la meme configuration.
_thinking_par_modele: dict[str, types.ThinkingConfig | None] = {}


def _is_invalid_argument(exc: Exception) -> bool:
    message = str(exc)
    return "INVALID_ARGUMENT" in message or "400" in message[:40]


def _est_definitif(exc: Exception) -> bool:
    """Erreur qu'une nouvelle tentative ne resoudra pas.

    Quota epuise, cle refusee, modele inexistant : insister ne fait que
    multiplier les appels — et, sur un quota, aggraver la situation.
    """
    haut = str(exc).upper()
    return any(
        marqueur in haut
        for marqueur in ("RESOURCE_EXHAUSTED", "UNAUTHENTICATED", "PERMISSION_DENIED", "NOT_FOUND")
    )


def _log_usage(response: types.GenerateContentResponse, model: str, elapsed: float) -> dict[str, int]:
    usage = getattr(response, "usage_metadata", None)
    stats = {
        "prompt": getattr(usage, "prompt_token_count", 0) or 0,
        "reponse": getattr(usage, "candidates_token_count", 0) or 0,
        "raisonnement": getattr(usage, "thoughts_token_count", 0) or 0,
        "total": getattr(usage, "total_token_count", 0) or 0,
    }
    logger.info(
        "Gemini %s — %.2fs | tokens prompt=%d reponse=%d raisonnement=%d total=%d",
        model,
        elapsed,
        stats["prompt"],
        stats["reponse"],
        stats["raisonnement"],
        stats["total"],
    )
    return stats


def _erreur_lisible(exc: Exception) -> ExtractionError:
    message = str(exc)
    haut = message.upper()
    if "API_KEY" in haut or "API KEY" in haut or "UNAUTHENTICATED" in haut or "PERMISSION_DENIED" in haut:
        return ExtractionError("Cle API Gemini refusee. Verifie GEMINI_API_KEY.", status_code=502)
    if "RESOURCE_EXHAUSTED" in haut or "429" in message[:40]:
        modele = get_settings().gemini_model
        if "PERDAY" in haut.replace("_", "") or "FREETIER" in haut.replace("_", "").replace("-", ""):
            # Palier gratuit : 20 requetes/jour/modele. Les alias « latest »
            # partagent le compteur du modele qu'ils designent.
            return ExtractionError(
                f"Quota gratuit épuisé pour « {modele} » (20 requêtes par jour et par modèle). "
                "Activez la facturation sur le projet Google AI Studio, ou changez GEMINI_MODEL "
                "dans backend/.env pour un modèle dont le quota reste disponible.",
                status_code=502,
            )
        return ExtractionError(
            "Trop de requêtes envoyées coup sur coup. Réessaie dans quelques instants.",
            status_code=502,
        )
    if "NOT_FOUND" in haut or "404" in message[:40]:
        return ExtractionError(
            "Le modele configure n'est pas disponible pour cette cle. "
            "Ajuste GEMINI_MODEL dans backend/.env.",
            status_code=502,
        )
    return ExtractionError("Le service d'analyse est momentanement indisponible.", status_code=502)


def _appeler_modele(
    client: genai.Client, model: str, contents: list[types.Content]
) -> tuple[types.GenerateContentResponse, float]:
    """Appelle Gemini en negociant la configuration de raisonnement acceptee.

    La variante retenue est memorisee par modele pour ne pas repayer la
    detection a chaque requete, mais elle est reevaluee des qu'un appel est
    rejete pour argument invalide : le modele a pu changer sous l'alias.
    """
    memorisee = model in _thinking_par_modele
    variantes = [_thinking_par_modele[model]] if memorisee else _thinking_variants()
    derniere_erreur: Exception | None = None

    for thinking in variantes:
        config = types.GenerateContentConfig(
            response_mime_type="application/json",
            response_schema=RESPONSE_SCHEMA,
            temperature=0.0,
            # Tache simple et cadree : on bride le raisonnement pour limiter la facture.
            thinking_config=thinking,
        )
        for tentative in range(2):
            started = time.perf_counter()
            try:
                response = client.models.generate_content(model=model, contents=contents, config=config)
                if not memorisee:
                    _thinking_par_modele[model] = thinking
                    logger.info("Configuration de raisonnement retenue pour %s : %s", model, thinking)
                return response, time.perf_counter() - started
            except Exception as exc:
                derniere_erreur = exc
                if _est_definitif(exc):
                    # Inutile de reessayer ni de changer de variante.
                    raise _erreur_lisible(exc) from exc
                if _is_invalid_argument(exc):
                    if memorisee:
                        # La variante memorisee n'est plus acceptee : on renegocie.
                        logger.warning("Variante memorisee refusee par %s, renegociation", model)
                        _thinking_par_modele.pop(model, None)
                        return _appeler_modele(client, model, contents)
                    logger.info("Variante de raisonnement refusee par %s, essai suivant", model)
                    break  # variante inadaptee au modele : on passe a la suivante
                logger.warning("Appel Gemini en echec (tentative %d) : %s", tentative + 1, exc)
                if tentative == 0:
                    time.sleep(1.2)

    raise _erreur_lisible(derniere_erreur or RuntimeError("appel Gemini impossible"))


def ping_model() -> dict[str, object]:
    """Appel minimal au modele configure, pour diagnostiquer une panne d'acces."""
    settings = get_settings()
    depart = time.perf_counter()
    try:
        client = get_client()
        reponse = client.models.generate_content(
            model=settings.gemini_model,
            contents="ping",
            config=types.GenerateContentConfig(max_output_tokens=8),
        )
    except ExtractionError as exc:
        return {"ok": False, "model": settings.gemini_model, "erreur": exc.message}
    except Exception as exc:
        message = str(exc)
        logger.warning("Sonde modele en echec pour %s : %s", settings.gemini_model, message)
        return {
            "ok": False,
            "model": settings.gemini_model,
            "erreur": _erreur_lisible(exc).message,
            "detail": message[:400],
        }

    usage = getattr(reponse, "usage_metadata", None)
    return {
        "ok": True,
        "model": settings.gemini_model,
        "latenceMs": round((time.perf_counter() - depart) * 1000),
        "tokens": getattr(usage, "total_token_count", 0) or 0,
    }


def extract_supplies(filename: str, payload: bytes) -> tuple[Extraction, dict[str, int]]:
    """Envoie le document a Gemini et renvoie l'extraction validee + les tokens."""
    settings = get_settings()
    document = build_document_parts(filename, payload)
    client = get_client()

    contents = [types.Content(role="user", parts=[types.Part.from_text(text=PROMPT), *document.parts])]
    response, duree = _appeler_modele(client, settings.gemini_model, contents)

    tokens = _log_usage(response, settings.gemini_model, duree)
    data = _parse_payload(response.text)

    try:
        extraction = Extraction.model_validate(data)
    except Exception as exc:
        logger.error("Donnees Gemini invalides : %s", str(data)[:400])
        raise ExtractionError("Les donnees extraites sont incompletes. Reessaie.") from exc

    # Filet de securite : on ecarte les lignes vides et les classes sans article.
    extraction.elaguer()

    logger.info(
        "Extraction %s (%s) — %d classe(s) %s, %d article(s), rubriques %s",
        filename,
        document.source,
        len(extraction.listes),
        extraction.classes or ["(non identifiee)"],
        len(extraction.articles),
        list(extraction.repartition),
    )
    return extraction, tokens
