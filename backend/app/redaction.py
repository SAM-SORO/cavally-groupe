"""Mise en document d'une liste saisie au clavier.

Un client qui n'a pas de fichier tape sa liste directement dans l'interface.
L'entreprise, elle, doit toujours recevoir un **document** exploitable sur
WhatsApp — jamais un long message texte qui se perd dans une conversation.
Ce module fait la conversion : texte brut -> `.docx`.

Rien d'autre. **Aucune extraction, aucune analyse** : les lignes sont reprises
telles quelles. C'est l'outil interne qui, plus tard et de son cote, passera ce
document dans Gemini.
"""

from __future__ import annotations

import logging
import re
import unicodedata
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .models import Client

logger = logging.getLogger("cavally.redaction")

MIME_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

LONGUEUR_MIN_TEXTE = 3
LONGUEUR_MAX_TEXTE = 20_000

# Charte : jaune de la marque, noir. Rien d'autre.
JAUNE = "FFB800"
NOIR = RGBColor(0x00, 0x00, 0x00)
GRIS = RGBColor(0x66, 0x66, 0x66)

# Puces decoratives en tete de ligne. Les NOMBRES ne sont volontairement pas
# retires : « 3 cahiers » comme « 1. Cahier » peuvent porter une quantite, et
# la perdre fausserait le devis.
_PUCE = re.compile(r"^[-–—*•·]+\s*")


def _lignes_utiles(texte: str) -> list[str]:
    """Lignes non vides, debarrassees de leur puce decorative."""
    lignes = []
    for brute in texte.splitlines():
        ligne = _PUCE.sub("", brute.strip()).strip()
        if ligne:
            lignes.append(ligne)
    return lignes


def _ombrer(paragraphe, couleur: str) -> None:
    """Aplat de couleur derriere un paragraphe (python-docx ne l'expose pas)."""
    fond = OxmlElement("w:shd")
    fond.set(qn("w:val"), "clear")
    fond.set(qn("w:fill"), couleur)
    paragraphe._p.get_or_add_pPr().append(fond)


def nom_fichier(client: Client, horodatage: datetime | None = None) -> str:
    """« demande-koffi-aya-20260726-1652.docx » — ASCII, sans espace."""
    horodatage = horodatage or datetime.now()
    sans_accent = unicodedata.normalize("NFKD", client.nom_complet or "client")
    sans_accent = sans_accent.encode("ascii", "ignore").decode("ascii")
    fragment = re.sub(r"[^a-zA-Z0-9]+", "-", sans_accent).strip("-").lower() or "client"
    return f"demande-{fragment[:40]}-{horodatage:%Y%m%d-%H%M}.docx"


def construire_docx(client: Client, texte: str, horodatage: datetime | None = None) -> bytes:
    """Document de demande : en-tete d'identification, puis la liste."""
    horodatage = horodatage or datetime.now()
    lignes = _lignes_utiles(texte)

    document = Document()

    section = document.sections[0]
    section.left_margin = section.right_margin = Cm(2.2)
    section.top_margin = section.bottom_margin = Cm(2)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # — Titre —
    titre = document.add_paragraph()
    trait = titre.add_run("DEMANDE DE DEVIS")
    trait.bold = True
    trait.font.size = Pt(20)
    trait.font.color.rgb = NOIR
    titre.paragraph_format.space_after = Pt(2)

    marque = document.add_paragraph()
    ligne_marque = marque.add_run("Cavally Livres — fournitures scolaires")
    ligne_marque.font.size = Pt(10)
    ligne_marque.font.color.rgb = GRIS
    marque.paragraph_format.space_after = Pt(14)

    # — Identite du demandeur —
    # C'est le coeur du document pour l'equipe : savoir qui rappeler.
    bandeau = document.add_paragraph()
    _ombrer(bandeau, JAUNE)
    intitule = bandeau.add_run("  SOUMISSIONNAIRE")
    intitule.bold = True
    intitule.font.size = Pt(10)
    intitule.font.color.rgb = NOIR
    bandeau.paragraph_format.space_after = Pt(8)

    champs = [
        ("Nom", client.nom_complet),
        ("Téléphone", client.contact or "non renseigné"),
        ("Email", client.email),
    ]
    if client.etablissement:
        champs.append(("Établissement", client.etablissement))
    champs.append(("Déposée le", f"{horodatage:%d/%m/%Y à %H:%M}"))

    tableau = document.add_table(rows=0, cols=2)
    tableau.style = "Table Grid"
    for libelle, valeur in champs:
        cellules = tableau.add_row().cells
        cellules[0].width = Cm(4.2)
        cellules[1].width = Cm(12)
        gauche = cellules[0].paragraphs[0].add_run(libelle)
        gauche.bold = True
        gauche.font.size = Pt(10)
        droite = cellules[1].paragraphs[0].add_run(str(valeur))
        droite.font.size = Pt(10)

    document.add_paragraph().paragraph_format.space_after = Pt(6)

    # — La liste, telle que le client l'a ecrite —
    entete = document.add_paragraph()
    _ombrer(entete, JAUNE)
    intitule_liste = entete.add_run("  LISTE DE FOURNITURES")
    intitule_liste.bold = True
    intitule_liste.font.size = Pt(10)
    intitule_liste.font.color.rgb = NOIR
    entete.paragraph_format.space_after = Pt(8)

    if lignes:
        for ligne in lignes:
            element = document.add_paragraph(ligne, style="List Bullet")
            element.paragraph_format.space_after = Pt(2)
    else:
        vide = document.add_paragraph()
        mention = vide.add_run("(aucune ligne exploitable dans la saisie)")
        mention.italic = True
        mention.font.color.rgb = GRIS

    # — Pied —
    pied = document.add_paragraph()
    pied.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pied.paragraph_format.space_before = Pt(18)
    note = pied.add_run("Liste saisie par le client sur la plateforme Cavally Livres.")
    note.font.size = Pt(8)
    note.font.color.rgb = GRIS

    tampon = BytesIO()
    document.save(tampon)
    contenu = tampon.getvalue()

    logger.info(
        "Document rédigé pour #%d — %d ligne(s), %.1f Ko",
        client.id,
        len(lignes),
        len(contenu) / 1024,
    )
    return contenu
